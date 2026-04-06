import os

import docx
import PyPDF2

from .constants import ALLOWED_EXTENSIONS


class UnsupportedFileFormatError(ValueError):
    pass


def file_extension(filename):
    if not filename:
        return ""
    return os.path.splitext(filename.lower())[1]


def validate_upload_extension(filename):
    ext = file_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileFormatError("Unsupported file format")


def _docx_all_text(document):
    """Paragraphs plus table cell text (common in lecture notes)."""
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    chunks.append(t)
    return "\n".join(chunks)


def extract_text_from_file(file):
    """Read plain text from an uploaded file. Resets read position when possible."""
    if hasattr(file, "seek"):
        file.seek(0)

    filename = (getattr(file, "name", "") or "").lower()
    ext = file_extension(filename)

    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileFormatError("Unsupported file format")

    try:
        if ext == ".pdf":
            reader = PyPDF2.PdfReader(file)
            if getattr(reader, "is_encrypted", False):
                try:
                    ok = reader.decrypt("")
                except Exception:
                    ok = 0
                if ok == 0:
                    raise ValueError(
                        "PDF is password-protected; upload an unlocked copy."
                    )
            text = "".join(page.extract_text() or "" for page in reader.pages)
        elif ext == ".docx":
            document = docx.Document(file)
            text = _docx_all_text(document)
        else:
            raw = file.read()
            text = raw.decode("utf-8-sig", errors="replace")
    except UnsupportedFileFormatError:
        raise
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not extract text: {exc}") from exc
    finally:
        if hasattr(file, "seek"):
            file.seek(0)

    return text.strip()
